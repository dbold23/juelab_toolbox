#!/usr/bin/env python3
"""Convert full_paper.md to a formatted Word document with embedded figures."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

PAPER_DIR = Path(__file__).parent
MD_PATH = PAPER_DIR / "full_paper.md"
FIGURES_DIR = PAPER_DIR / "figures"
OUT_PATH = PAPER_DIR / "full_paper.docx"


def setup_styles(doc):
    """Configure document styles for academic paper."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 2.0

    for level, size in [(1, 14), (2, 13), (3, 12)]:
        name = f'Heading {level}'
        s = doc.styles[name]
        s.font.name = 'Times New Roman'
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)


def add_formatted_paragraph(doc, text, style='Normal', bold=False):
    """Add a paragraph with inline bold/italic markdown formatting."""
    p = doc.add_paragraph(style=style)
    # Split on **bold** and *italic* patterns
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


def convert_md_to_docx(md_path, out_path):
    doc = Document()
    setup_styles(doc)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    text = md_path.read_text(encoding='utf-8')
    lines = text.split('\n')

    # Track figure insertion points
    figure_map = {
        'Figure 1': 'figure1_pipeline_schematic.png',
        'Figure 2': 'figure2_dataset_overview.png',
        'Figure 3': 'figure3_synthetic_validation.png',
        'Figure 4': 'figure4_truncation_comparison.png',
        'Figure 5': 'figure5_haldane_vs_gompertz.png',
        'Figure 6': 'figure6_bayesian_ki_forest.png',
        'Figure 7': 'figure7_inter_operator.png',
        'Figure 8': 'figure8_representative_fits.png',
    }

    # Section tracking for figure insertion
    current_section = ""
    figure_inserted = set()

    # Figure-to-section mapping (insert figure after first mention in results)
    figure_sections = {
        'Figure 1': '2.2 Pipeline Architecture',
        'Figure 2': '3.1 Dataset Overview',
        'Figure 3': '3.3 Synthetic Validation',
        'Figure 4': '3.4 Truncation Method Comparison',
        'Figure 5': '3.5 Haldane versus Gompertz',
        'Figure 6': '3.6 Bayesian Inhibition Constant',
        'Figure 7': '3.7 Inter-Operator Reproducibility',
        'Figure 8': 'Figure Legends',
    }

    i = 0
    in_equation_block = False
    skip_hr = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip horizontal rules
        if stripped == '---':
            i += 1
            continue

        # Title (# level 1 at very start)
        if stripped.startswith('# ') and not stripped.startswith('## '):
            heading_text = stripped[2:].strip()
            level = 1

            # Main title gets special formatting
            if 'automated computational pipeline' in heading_text.lower():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(heading_text)
                run.bold = True
                run.font.size = Pt(16)
                run.font.name = 'Times New Roman'
                i += 1
                continue

            doc.add_heading(heading_text, level=level)
            i += 1
            continue

        # H2 headings
        if stripped.startswith('## '):
            heading_text = stripped[3:].strip()
            doc.add_heading(heading_text, level=2)
            current_section = heading_text
            i += 1
            continue

        # H3 headings
        if stripped.startswith('### '):
            heading_text = stripped[4:].strip()
            doc.add_heading(heading_text, level=3)
            current_section = heading_text
            i += 1
            continue

        # Empty lines
        if not stripped:
            i += 1
            continue

        # LaTeX display equations ($$...$$)
        if stripped.startswith('$$'):
            # Collect until closing $$
            eq_lines = [stripped[2:]]
            if stripped.endswith('$$') and len(stripped) > 4:
                eq_text = stripped[2:-2].strip()
            else:
                i += 1
                while i < len(lines):
                    eqline = lines[i].strip()
                    if eqline.endswith('$$'):
                        eq_lines.append(eqline[:-2])
                        break
                    eq_lines.append(eqline)
                    i += 1
                eq_text = ' '.join(eq_lines).strip()

            # Clean up LaTeX for Word (basic conversion)
            eq_text = eq_text.replace(r'\cdot', '*')
            eq_text = eq_text.replace(r'\exp', 'exp')
            eq_text = eq_text.replace(r'\max', 'max')
            eq_text = eq_text.replace(r'\ln', 'ln')
            eq_text = eq_text.replace(r'\sqrt', 'sqrt')
            eq_text = eq_text.replace(r'\frac', 'frac')
            eq_text = eq_text.replace(r'\left(', '(')
            eq_text = eq_text.replace(r'\right)', ')')
            eq_text = eq_text.replace(r'\left)', ')')
            eq_text = eq_text.replace(r'\Big(', '(')
            eq_text = eq_text.replace(r'\Big)', ')')
            eq_text = eq_text.replace(r'\Bigg(', '(')
            eq_text = eq_text.replace(r'\Bigg)', ')')
            eq_text = eq_text.replace(r'\bigg(', '(')
            eq_text = eq_text.replace(r'\bigg)', ')')
            eq_text = eq_text.replace(r'\!', '')
            eq_text = eq_text.replace(r'\;', ' ')
            eq_text = eq_text.replace(r'\,', ' ')
            eq_text = eq_text.replace(r'\text{', '')
            eq_text = eq_text.replace(r'\overline{', '')
            eq_text = re.sub(r'\\hat\{(\w)\}', r'\1_hat', eq_text)
            eq_text = re.sub(r'\\bar\{(\w)\}', r'\1_bar', eq_text)
            eq_text = eq_text.replace('}', '')
            eq_text = eq_text.replace('{', '')

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(eq_text)
            run.font.name = 'Cambria Math'
            run.font.size = Pt(11)
            run.italic = True
            i += 1
            continue

        # Author/affiliation lines (centered)
        if any(x in stripped for x in ['Daniel Sambold', 'Department of Marine Science']):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if 'Daniel' in stripped:
                run.font.size = Pt(13)
            i += 1
            continue

        # Regular paragraph
        # Clean up inline LaTeX
        clean = stripped
        clean = clean.replace('$R^2$', 'R²')
        clean = clean.replace('$\\hat{R}$', 'R-hat')
        clean = clean.replace('$R^{2}$', 'R²')
        clean = re.sub(r'\$([^$]+)\$', lambda m: m.group(1)
                       .replace(r'\mu_m', 'μ_m')
                       .replace(r'\mu', 'μ')
                       .replace(r'\lambda', 'λ')
                       .replace(r'\sigma', 'σ')
                       .replace(r'\Delta', 'Δ')
                       .replace(r'\times', '×')
                       .replace(r'\leq', '≤')
                       .replace(r'\geq', '≥')
                       .replace(r'\text{', '')
                       .replace(r'\log', 'log')
                       .replace(r'\sim', '~')
                       .replace('{', '').replace('}', '')
                       .replace('^2', '²')
                       .replace('^{-1}', '⁻¹')
                       .replace('_max', '_max')
                       .replace('_obs', '_obs'),
                       clean)
        clean = clean.replace('`', '')

        add_formatted_paragraph(doc, clean)

        # Check if we should insert a figure after this paragraph
        for fig_name, section_trigger in figure_sections.items():
            if (fig_name not in figure_inserted and
                section_trigger in current_section and
                fig_name in figure_map):
                fig_file = FIGURES_DIR / figure_map[fig_name]
                if fig_file.exists():
                    doc.add_paragraph()  # spacing
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(fig_file), width=Inches(6.0))
                    # Caption
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap.add_run(f'{fig_name}')
                    cap_run.bold = True
                    cap_run.font.size = Pt(10)
                    figure_inserted.add(fig_name)

        i += 1

    doc.save(str(out_path))
    print(f"Saved: {out_path}")
    print(f"Size: {out_path.stat().st_size / 1024:.0f} KB")


if __name__ == '__main__':
    convert_md_to_docx(MD_PATH, OUT_PATH)
